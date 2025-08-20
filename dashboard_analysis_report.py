import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from pathlib import Path
import sys

def load_and_analyze_data():
    """Load and analyze the dataset to extract key insights"""
    try:
        # Try to load the CSV file
        csv_paths = [
            Path("ai_vs_human_content_dataset.csv"),
            Path(r"C:\Users\Sakshi\Desktop\ai_vs_human_content_dataset.csv"),
        ]
        
        csv_path = None
        for path in csv_paths:
            if path.exists():
                csv_path = path
                break
                
        if csv_path is None:
            raise FileNotFoundError("Could not find the CSV file")
            
        df = pd.read_csv(csv_path)
        
        # Clean data (same as dashboard)
        if "Date" in df.columns:
            df["Date"] = pd.to_datetime(df["Date"], format="%d-%m-%Y", errors="coerce")
        
        if "Author_Type" in df.columns:
            df["Author_Type"] = df["Author_Type"].astype(str).str.strip().str.title()
        
        if set(["Likes", "Comments", "Shares"]).issubset(df.columns):
            df["Engagement_Score"] = (
                df["Likes"].fillna(0) + 
                df["Comments"].fillna(0) * 2 + 
                df["Shares"].fillna(0) * 3
            )
        
        df = df.drop_duplicates().reset_index(drop=True)
        
        # Calculate key metrics
        insights = {}
        
        # Overall metrics
        insights['total_posts'] = len(df)
        insights['ai_posts'] = len(df[df['Author_Type'] == 'Ai'])
        insights['human_posts'] = len(df[df['Author_Type'] == 'Human'])
        
        # Engagement metrics by author type
        ai_data = df[df['Author_Type'] == 'Ai']
        human_data = df[df['Author_Type'] == 'Human']
        
        insights['ai_avg_engagement'] = ai_data['Engagement_Score'].mean()
        insights['human_avg_engagement'] = human_data['Engagement_Score'].mean()
        insights['ai_avg_likes'] = ai_data['Likes'].mean()
        insights['human_avg_likes'] = human_data['Likes'].mean()
        insights['ai_avg_comments'] = ai_data['Comments'].mean()
        insights['human_avg_comments'] = human_data['Comments'].mean()
        insights['ai_avg_shares'] = ai_data['Shares'].mean()
        insights['human_avg_shares'] = human_data['Shares'].mean()
        
        # Content type analysis
        content_performance = df.groupby(['Content_Type', 'Author_Type'])['Engagement_Score'].mean().unstack()
        insights['content_performance'] = content_performance
        
        # Top performing content types
        top_content = df.groupby('Content_Type')['Engagement_Score'].mean().sort_values(ascending=False)
        insights['top_content_types'] = top_content.head(3)
        
        # Engagement trend
        if 'Date' in df.columns:
            monthly_trend = df.groupby([df['Date'].dt.to_period('M'), 'Author_Type'])['Engagement_Score'].mean().unstack()
            insights['monthly_trend'] = monthly_trend
        
        return insights, df
        
    except Exception as e:
        print(f"Error loading data: {e}")
        return None, None

def create_analysis_report(insights, df):
    """Create a Word document with the analysis report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('AI vs Human Content Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run('This report analyzes the performance of AI-generated content versus human-written content based on social media engagement metrics. ')
    summary.add_run(f'We examined {insights["total_posts"]} posts ({insights["ai_posts"]} AI-generated, {insights["human_posts"]} human-written) ')
    summary.add_run('to understand which type of content drives better audience engagement.')
    
    # Key Findings
    doc.add_heading('Key Findings', level=1)
    
    # Finding 1: Overall Engagement Comparison
    doc.add_heading('1. AI vs Human Engagement Performance', level=2)
    
    ai_engagement = insights['ai_avg_engagement']
    human_engagement = insights['human_avg_engagement']
    
    if ai_engagement > human_engagement:
        diff_percent = ((ai_engagement - human_engagement) / human_engagement) * 100
        finding1 = doc.add_paragraph()
        finding1.add_run(f'AI-generated content outperformed human-written content by {diff_percent:.1f}% in overall engagement. ')
        finding1.add_run(f'AI posts achieved an average engagement score of {ai_engagement:.0f}, ')
        finding1.add_run(f'while human posts averaged {human_engagement:.0f}.')
    else:
        diff_percent = ((human_engagement - ai_engagement) / ai_engagement) * 100
        finding1 = doc.add_paragraph()
        finding1.add_run(f'Human-written content outperformed AI-generated content by {diff_percent:.1f}% in overall engagement. ')
        finding1.add_run(f'Human posts achieved an average engagement score of {human_engagement:.0f}, ')
        finding1.add_run(f'while AI posts averaged {ai_engagement:.0f}.')
    
    # Finding 2: Engagement Breakdown
    doc.add_heading('2. Engagement Metric Breakdown', level=2)
    
    breakdown = doc.add_paragraph()
    breakdown.add_run('When we break down engagement by individual metrics:\n\n')
    
    # Likes comparison
    ai_likes = insights['ai_avg_likes']
    human_likes = insights['human_avg_likes']
    breakdown.add_run(f'• Likes: AI posts received {ai_likes:.0f} average likes vs {human_likes:.0f} for human posts\n')
    
    # Comments comparison
    ai_comments = insights['ai_avg_comments']
    human_comments = insights['human_avg_comments']
    breakdown.add_run(f'• Comments: AI posts received {ai_comments:.0f} average comments vs {human_comments:.0f} for human posts\n')
    
    # Shares comparison
    ai_shares = insights['ai_avg_shares']
    human_shares = insights['human_avg_shares']
    breakdown.add_run(f'• Shares: AI posts received {ai_shares:.0f} average shares vs {human_shares:.0f} for human posts')
    
    # Finding 3: Content Type Performance
    doc.add_heading('3. Content Type Analysis', level=2)
    
    content_analysis = doc.add_paragraph()
    content_analysis.add_run('Different content types perform differently for AI vs Human authors:\n\n')
    
    # Get top performing content types
    top_content = insights['top_content_types']
    content_analysis.add_run('Top 3 performing content types overall:\n')
    for i, (content_type, score) in enumerate(top_content.items(), 1):
        content_analysis.add_run(f'{i}. {content_type}: {score:.0f} average engagement\n')
    
    # Finding 4: Business Implications
    doc.add_heading('4. Business Implications', level=2)
    
    implications = doc.add_paragraph()
    implications.add_run('Based on our analysis, here are the key business implications:\n\n')
    
    if ai_engagement > human_engagement:
        implications.add_run('• AI-generated content shows strong potential for driving engagement\n')
        implications.add_run('• Consider incorporating AI tools into content creation workflows\n')
        implications.add_run('• AI content may be more cost-effective for maintaining consistent posting schedules\n')
    else:
        implications.add_run('• Human-written content maintains an edge in audience connection\n')
        implications.add_run('• Focus on human creativity and personal touch in content strategy\n')
        implications.add_run('• Consider AI as a supplement rather than replacement for human content\n')
    
    implications.add_run('• Monitor engagement trends over time to adapt strategy accordingly\n')
    implications.add_run('• Test different content types to optimize for your specific audience')
    
    # Recommendations
    doc.add_heading('Recommendations', level=1)
    
    recommendations = doc.add_paragraph()
    recommendations.add_run('Based on our findings, we recommend:\n\n')
    
    recommendations.add_run('1. ')
    if ai_engagement > human_engagement:
        recommendations.add_run('Gradually incorporate AI-generated content into your content strategy\n')
    else:
        recommendations.add_run('Maintain focus on human-created content while exploring AI as a supplement\n')
    
    recommendations.add_run('2. ')
    recommendations.add_run('Focus on content types that show the highest engagement for your target audience\n')
    
    recommendations.add_run('3. ')
    recommendations.add_run('Implement A/B testing to compare AI vs Human content performance in your specific context\n')
    
    recommendations.add_run('4. ')
    recommendations.add_run('Monitor engagement metrics regularly to track performance trends\n')
    
    recommendations.add_run('5. ')
    recommendations.add_run('Consider hybrid approaches: use AI for initial drafts and human editors for final touches')
    
    # Methodology
    doc.add_heading('Methodology', level=1)
    
    methodology = doc.add_paragraph()
    methodology.add_run('This analysis was conducted using:\n\n')
    methodology.add_run(f'• Dataset: {insights["total_posts"]} social media posts\n')
    methodology.add_run(f'• Time period: {df["Date"].min().strftime("%B %Y")} to {df["Date"].max().strftime("%B %Y")}\n')
    methodology.add_run('• Engagement Score Formula: Likes + (2 × Comments) + (3 × Shares)\n')
    methodology.add_run('• Analysis tools: Python, Pandas, and statistical analysis\n')
    methodology.add_run('• Visualization: Interactive dashboard with real-time filtering capabilities')
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run('This analysis provides valuable insights into the performance of AI-generated versus human-written content. ')
    
    if ai_engagement > human_engagement:
        conclusion.add_run('The data suggests that AI-generated content can effectively drive engagement, ')
        conclusion.add_run('offering opportunities for cost-effective content creation while maintaining audience interest.')
    else:
        conclusion.add_run('The data reinforces the value of human creativity in content creation, ')
        conclusion.add_run('while also highlighting opportunities for strategic use of AI tools.')
    
    conclusion.add_run('\n\nHowever, the most successful content strategies will likely involve a thoughtful blend of both approaches, ')
    conclusion.add_run('tailored to your specific audience and business objectives.')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Report generated on: ' + pd.Timestamp.now().strftime('%B %d, %Y'))
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def main():
    print("Loading and analyzing data...")
    insights, df = load_and_analyze_data()
    
    if insights is None:
        print("Failed to load data. Please ensure the CSV file is available.")
        return
    
    print("Creating analysis report...")
    doc = create_analysis_report(insights, df)
    
    # Save the document
    output_path = "AI_vs_Human_Content_Analysis_Report.docx"
    doc.save(output_path)
    
    print(f"Analysis report saved as: {output_path}")
    print("\nKey insights from the analysis:")
    print(f"- Total posts analyzed: {insights['total_posts']}")
    print(f"- AI posts: {insights['ai_posts']}, Human posts: {insights['human_posts']}")
    print(f"- AI avg engagement: {insights['ai_avg_engagement']:.0f}")
    print(f"- Human avg engagement: {insights['human_avg_engagement']:.0f}")

if __name__ == "__main__":
    main()
