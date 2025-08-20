import sys
from pathlib import Path


def main() -> int:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover
        print("python-docx is not installed. Run: python -m pip install python-docx", file=sys.stderr)
        print(f"Details: {exc}", file=sys.stderr)
        return 1

    # Handle both possible filenames with and without the emoji, just in case
    candidates = [
        Path("\ud83d\udd25 Project.docx"),  # literal escape for the emoji
        Path("🔥 Project.docx"),
        Path("Project.docx"),
    ]

    doc_path = None
    for p in candidates:
        if p.exists():
            doc_path = p
            break

    if doc_path is None:
        print("Could not find the project document. Looked for: '🔥 Project.docx' and 'Project.docx'", file=sys.stderr)
        return 1

    doc = Document(str(doc_path))

    # Persist content to a UTF-8 text file to avoid console encoding issues
    out_path = Path("project_doc.txt")
    with out_path.open("w", encoding="utf-8") as fh:
        fh.write("=== Document Text ===\n")
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                fh.write(text + "\n")

        # Extract table contents if present
        if doc.tables:
            fh.write("\n=== Document Tables ===\n")
            for ti, table in enumerate(doc.tables, start=1):
                fh.write(f"-- Table {ti} --\n")
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    fh.write(" | ".join(cells) + "\n")

    print(f"Wrote extracted text to {out_path}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())


